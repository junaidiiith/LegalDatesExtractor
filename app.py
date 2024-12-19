import time
import PyPDF2
import streamlit as st

import openai
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from stqdm import stqdm


model = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
system_prompt = {
    "role": "system", 
    "content": "You are an expert in reading legal criminal cases and you need to extract the dates and the details about what happened on that date."
}


def get_dates_details(text, messages=None):
	if messages is None:
		messages = [system_prompt]

	prompt = \
	f"""
	You need to extract the dates and the details about what happened on that date. 
	Provide empty string in case there are no dates present.
	Any date that involves an event should be extracted.

	Extract the dates from the following text:
	{text}

	Example: 
	Input: The present reference is an abuse of the process. 
	The claimants, on an earlier occasion, had initiated another Arbitration reference, involving disputes purportedly arising out of the self same agreement dated 10th November, 2008, against the present respondents, which ultimately resulted into an Award dated 5th May, 2022.  

	Output:
	- 10th November, 2008: The claimants initiated an Arbitration reference involving disputes purportedly arising out of the self same agreement.
	- 5th May, 2022: The Arbitration reference resulted into an Award.
	"""
	response = model.chat.completions.create(
		model="gpt-4o",
		messages=messages + [
			{"role": "user", "content": prompt}
		],
	)
	ai_response = response.choices[0].message.content
	# time.sleep(0.01)
	# ai_response = "Good job! You have successfully extracted the dates and the details about what happened on that date."
	messages.append({"role": "assistant", "content": ai_response})
	
	return messages


def refine_dates_texts():
	doc_dates = st.session_state.get("doc_dates", None)

	if not doc_dates:
		st.error("No document dates found.")
		return
	
	doc_dates_refined = st.session_state.get("doc_dates_refined", None)
		
	if not doc_dates_refined:
		prompt = \
		f"""
		You are provided with a list of dates and what happened on those dates.
		The event on all these dates may be interconnected and therefore some date might be missing some information.
		You need to refine the the details about what happened on each date from the following text:
		{doc_dates}
		"""
		with st.spinner("Refining dates and details..."):
			# time.sleep(1)
			response = model.chat.completions.create(
				model="gpt-4o",
				messages=[
					{"role": "system", "content": "You are an expert in reading legal criminal cases and finding the missing details about the dates and events."},
					{"role": "user", "content": prompt},
				],
			)
			doc_dates_refined = response.choices[0].message.content
			# doc_dates_refined = "Amazing! You have successfully refined the details about what happened on each date."
			
		st.session_state["doc_dates_refined"] = doc_dates_refined


def get_document_dates():
	text = st.session_state.get("doc_text", None)
	if not text:
		st.error("No text found in the document.")
		return
	
	doc_dates = st.session_state.get("doc_dates", None)

	if not doc_dates:
		chunk_size = st.session_state.get("chunk_size", 5000)
		chunk_overlap = st.session_state.get("chunk_overlap", 50)
		text_splitter = RecursiveCharacterTextSplitter(
			chunk_size=chunk_size, 
			chunk_overlap=chunk_overlap
		)
		texts = text_splitter.split_text(text)
		
		messages = None

		for text in stqdm(texts, desc="Extracting dates and details"):
			messages = get_dates_details(text, messages=messages)
		
		doc_dates = "\n".join([m['content'] for m in messages if m['role'] == 'assistant'])
		st.session_state["doc_dates"] = doc_dates
	

def get_document_text():
	uploaded_doc = st.session_state.get("doc_file", None)
	if not uploaded_doc:
		return
	if uploaded_doc.name.endswith(".pdf"):
		text = pdf_to_markdown(uploaded_doc)
		st.markdown(f"### Uploaded document {uploaded_doc.name} Info")
		
		num_pages = len(text.split('\n'))
		st.write(f"Number of pages: {num_pages}")
		st.write(f"Number of words: {len(text.split())}")

		
	elif uploaded_doc.name.endswith(".docx"):
		doc = Document(uploaded_doc)
		text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
		
		st.markdown(f"### Uploaded document {uploaded_doc.name} Info")
		st.write(f"Number of paragraphs: {len(doc.paragraphs)}")
		st.write(f"Number of words: {len(' '.join([p.text for p in doc.paragraphs]).split())}")
		
	else:
		st.error("Invalid file format. Only .docx and .pdf files are supported.")
		return

	st.session_state["doc_text"] = text

def pdf_to_markdown(pdf_file):
    # Initialize the PDF reader
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    texts = list()
    
    # Iterate through all pages and extract text
    for page in pdf_reader.pages:
        part_text = page.extract_text()
        texts.append(part_text)
    
    return "\n".join(texts)



def reset():
	st.session_state["doc_text"] = None
	st.session_state["doc_dates"] = None
	st.session_state["doc_dates_refined"] = None
	

def main():
	st.title("Legal Document Dates Extractor")
	st.file_uploader("Upload a document", type=["docx", "pdf"], key="doc_file")
	doc = st.session_state.get("doc_file", None)
	if doc:
		get_document_text()
		cols = st.columns(2)
		cols[0].number_input("Chunk size", value=5000, key="chunk_size", help="The size of the text to be processed at a time. Default is 5000.")
		cols[1].number_input("Chunk overlap", value=50, key="chunk_overlap", help="The overlap between chunks. Default is 50.")
		
		
		cols = st.columns(3)
		
		extract_dates = cols[0].button("Extract dates and details", key="extract_dates", on_click=get_document_dates)
		refine_dates = cols[1].button("Refine dates and details", key="refine_dates", on_click=refine_dates_texts)
		cols[2].button("Reset Document", key="reset", on_click=reset)

		if extract_dates:
			with st.container():
				doc_dates = st.session_state.get("doc_dates")
				st.markdown("### Extracted dates and details\n" + doc_dates)
				st.session_state["doc_dates"] = doc_dates
		
		if refine_dates:
			with st.container():
				doc_dates_refined = st.session_state.get("doc_dates_refined")
				st.markdown("### Refined dates and details")
				st.write(doc_dates_refined)

main()