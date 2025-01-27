import os
import streamlit as st
from docx import Document as DocxDocument
from streamlit.runtime.uploaded_file_manager import UploadedFile
from llama_parse import LlamaParse
from llama_index.core.schema import Document
from settings import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)
from llama_index.core.node_parser import SentenceSplitter
import uuid

def create_temp_file(upload_file: UploadedFile):
    uuid_str = str(uuid.uuid4())
    with open(f'{uuid_str}.pdf', 'wb') as f:
        f.write(upload_file.getvalue())
    pdf_file = f'{uuid_str}.pdf'
    return pdf_file

def pdf_to_markdown(pdf_file: UploadedFile):
    temp_file = create_temp_file(pdf_file)
    parser = LlamaParse(
        api_key=st.secrets['LLAMA_PARSE_API_KEY'],
        result_type="markdown"
    )
    
    try:
        with st.spinner("Parsing document..."):
            documents = parser.load_data(temp_file)
    except Exception as e:
        st.error(f"An error occurred while processing the document: {e}")
    full_text = "\n".join([d.text for d in documents])
    
    os.remove(temp_file)
    
    return full_text


def get_document_text():
    uploaded_doc: UploadedFile = st.session_state.get("doc_file", None)
    if not uploaded_doc:
        return

    if uploaded_doc.name.endswith(".txt"):
        text = uploaded_doc.read().decode("utf-8")
        st.session_state["doc"] = {
            "name": uploaded_doc.name,
            "text": text,
            "details": f"Number of lines: {len(text.splitlines())}\nNumber of words: {len(text.split())}"
        }
    
    elif uploaded_doc.name.endswith(".pdf"):
        text = pdf_to_markdown(uploaded_doc)

        st.session_state["doc"] = {
            "name": uploaded_doc.name,
            "text": text,
            "details": f"Number of pages: {len(text.split('\n'))}\nNumber of words: {len(text.split())}"
        }
        
    elif uploaded_doc.name.endswith(".docx"):
        temp_file = create_temp_file(uploaded_doc)
        doc = DocxDocument(temp_file)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        st.session_state["doc"] = {
            "name": uploaded_doc.name,
            "text": text,
            "details": f"Number of paragraphs: {len(doc.paragraphs)}\nNumber of words: {len(' '.join([p.text for p in doc.paragraphs]).split())}"
        }
        os.remove(temp_file)
        
    else:
        st.error("Invalid file format. Only .docx and .pdf files are supported.")
        return


def print_document_details():
    doc_text = st.session_state.get("doc_text", None)
    if not doc_text:
        return

    st.markdown(f"### Uploaded document {doc_text['name']} Info")
    st.write(doc_text["details"])



def get_nodes_from_documents(data: str):
    splitter = SentenceSplitter(
        chunk_size=CHUNK_SIZE, 
        chunk_overlap=CHUNK_OVERLAP
    )
    docs = splitter.split_text(data)
    print("Total number of documents generated:", len(docs))
    return [
        Document(text=doc)
        for doc in docs
    ]